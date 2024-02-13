import json 

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ALIGN_VERTICAL
from variable import LOGO_VIATTECH, FOLDER_DOCX, FOLDER_RESUME



def create_template(filename):
    print("################################### JSON To DOCX #####################################")

    ########### JSON file ##################
    folder = FOLDER_RESUME + '/' + filename + '/'
    
    
    personal_information = json.load(open(folder + 'personalInformation.json'))['personal_information']
    areas_expertize = json.load(open(folder + 'expertizeArea.json'))['areas_expertize']
    dialectes = json.load(open(folder + 'dialecte.json'))['dialectes']
    education = json.load(open(folder + 'education.json'))['educations']
    jobs = json.load(open(folder + 'jobs.json'))['jobs']
    softskills = json.load(open(folder + 'softSkills.json'))['softskills']
    hardskills = json.load(open(folder + 'hardSkills.json'))['hardskills']
    standards = json.load(open(folder + 'standards.json'))['standards']
    software = json.load(open(folder + 'software.json'))['softwares']
    
    
    areas_expertize_name = [area_expertize['area_expertize'] for area_expertize in areas_expertize]
    software_name = ['\n'.join(sw['software']) for sw in software]
    dialecte_name = [dialecte['dialecte'] for dialecte in dialectes]
    
    ########### DOCX file ##################

    # Create a new Document
    doc = Document()

    # Add a header to the document (present on all pages)
    header = doc.sections[0].header

   # Add a table with two cells for the image and title
    table = header.add_table(rows=1, cols=2, width=Inches(6))
    cell_image = table.cell(0, 0)
    cell_title = table.cell(0, 1)
    
    # Set the border size for the entire table
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    for border in run.element.xpath('.//w:tcBorders'):
                        border.attrib['sz'] = str(Pt(2).pt)  # Set the border size to 2 points

    # Add the image to the left cell
    run_image = cell_image.paragraphs[0].add_run()
    run_image.add_picture(LOGO_VIATTECH, width=Inches(1.0))  # Replace 'your_image_path.jpg' with the actual image path

    # Add the title to the right cell
    run_title = cell_title.paragraphs[0].add_run("Dossier Technique")
    run_title.bold = True
    cell_title.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Add a line below the table
    line_paragraph = header.add_paragraph()
    line_paragraph.add_run().add_break()
    line_paragraph.add_run().add_break()
    line_paragraph.runs[0].underline = True
    line_paragraph.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # Adjust cell alignment for the title
    cell_title.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.BOTTOM
    cell_title.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    cell_title.vertical_anchor = WD_ALIGN_VERTICAL.BOTTOM
    
    # Add a centered cell with border and text
    centered_cell = doc.add_table(rows=1, cols=1).cell(0, 0)
    centered_cell.paragraphs[0].text = personal_information["role"]
    centered_cell.paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    centered_cell.paragraphs[0].runs[0].font.size = Pt(14)
    for paragraph in centered_cell.paragraphs:
        for run in paragraph.runs:
            for border in run.element.xpath('.//w:tcBorders'):
                border.attrib['sz'] = str(Pt(1).pt)  # Set the border size to 1 point

    # Add a title "Domaine de compétences"
    doc.add_heading('Domaine de compétences', level=1)
    
    # Add a table with 6 rows and 2 columns
    competences_table = doc.add_table(rows=6, cols=2)

    # Define the list of points for the left side of the table
    competences_points = ["Secteur", "Methodologies", "Domaine", "Normes", "Logiciels", "Langues"]
    competences_points_had = [areas_expertize_name, softskills, hardskills,
                              standards, software_name, dialecte_name]

    # Populate the left side of the table with the points
    for row_index in range(len(competences_points)):
        cell = competences_table.cell(row_index, 0)
        cell.text = competences_points[row_index]
        cell = competences_table.cell(row_index, 1)
        cell.text = "\n".join(competences_points_had[row_index])


    # Add a title "Formation"
    doc.add_heading('Formation', level=1)
        
    lengthEducation = len(education)
    schools_table = doc.add_table(rows=lengthEducation, cols=2)
    
    
    rows = range(0, lengthEducation)
    for row_index in rows:
        education_row = education[row_index]
        cell = schools_table.cell(row_index, 0)
        cell.text = education_row["end_year"]
        cell = schools_table.cell(row_index, 1)
        cell.text = ", ".join([education_row["school_name"], education_row["degree"]])
        

    # Add a title "Projets"
    doc.add_heading('Projets', level=1)

    # Iterate through each job in the JSON object
    for job in jobs:
        # Add company_name as the title to the left
        doc.add_heading(job["company_name"], level=2)

        # Add start_date and end_date to the right of the title
        date_text = f"{job['start_date']} - {job['end_date']}"
        doc.add_paragraph(date_text).alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        # Add person_role below the title
        doc.add_paragraph(job["person_role"])

        # Add specification below person_role
        doc.add_paragraph(job["specification"])

        # Add skills as a list below specification
        skills_list = doc.add_paragraph()
        for skill in job["skills"]:
            skills_list.add_run(f"• {skill}\n")

        # Add a blank line between jobs
        doc.add_paragraph()


    # Save the document
    doc.save(FOLDER_DOCX + '/DT_' + personal_information["firstname"] +'_' + personal_information["lastname"] + '.docx')
    print("Template créé et enregistré sous : ", FOLDER_DOCX + '/DT_' + personal_information["firstname"] +'_' + personal_information["lastname"] + '.docx')
